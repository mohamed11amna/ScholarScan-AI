%%writefile app.py
import streamlit as st
import torch
import os
import gdown
import pytesseract
import io
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from transformers import PegasusTokenizer, PegasusForConditionalGeneration
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# CONFIGURATION & SETUP
st.set_page_config(page_title="ScholarScan AI", page_icon="📑", layout="wide")

# Custom CSS for a better look
st.markdown("""
<style>
    .main {background-color: #f5f5f5;}
    .stButton>button {width: 100%; border-radius: 5px; height: 3em; background-color: #FF4B4B; color: white;}
    .reportview-container .main .block-container{padding-top: 2rem;}
</style>
""", unsafe_allow_html=True)

# MODEL LOADING FUNCTIONS
@st.cache_resource
def load_model():
    """Downloads model if missing, then loads it into memory."""
    model_folder = "pegasus_model_local"
    drive_id = "1w_gEvMi1oPYESUjYP-0Wf4kfJIThJxV8" # YOUR MODEL ID
    
    # Download if not exists
    if not os.path.exists(model_folder):
        with st.spinner("📥 Downloading Model from Drive (This happens only once)..."):
            gdown.download_folder(id=drive_id, output=model_folder, quiet=True)
            
    device = "cuda" if torch.cuda.is_available() else "cpu"
    
    try:
        tokenizer = PegasusTokenizer.from_pretrained(model_folder)
        model = PegasusForConditionalGeneration.from_pretrained(model_folder, torch_dtype=torch.float16).to(device)
        return tokenizer, model, device
    except Exception as e:
        st.error(f"Failed to load model: {e}")
        return None, None, None

# TEXT EXTRACTION FUNCTIONS
def extract_from_pdf(file_bytes):
    """Converts PDF to Images then uses OCR."""
    with open("temp.pdf", "wb") as f:
        f.write(file_bytes.getbuffer())
    
    images = convert_from_path("temp.pdf")
    text = ""
    
    # Progress bar for OCR
    ocr_bar = st.progress(0)
    status = st.empty()
    
    for i, img in enumerate(images):
        status.text(f"Scanning Page {i+1}/{len(images)}...")
        text += pytesseract.image_to_string(img) + "\n"
        ocr_bar.progress((i + 1) / len(images))
        
    status.empty()
    ocr_bar.empty()
    return text

def extract_from_docx(file_obj):
    """Reads text from Word documents."""
    doc = Document(file_obj)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_from_image(file_obj):
    """Reads text from uploaded Image."""
    image = Image.open(file_obj)
    return pytesseract.image_to_string(image)

# SUMMARIZATION LOGIC
def chunk_text(text, limit=150):
    words = text.split()
    chunks = []
    current = []
    for word in words:
        current.append(word)
        if len(current) >= limit:
            chunks.append(" ".join(current))
            current = []
    if current:
        chunks.append(" ".join(current))
    return chunks

def summarize_text(text, tokenizer, model, device, chunk_size, beam_size):
    chunks = chunk_text(text, chunk_size)
    summaries = []
    
    prog_bar = st.progress(0)
    
    # Batch processing for speed (Batch size 4)
    batch_size = 4
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        inputs = tokenizer(batch, truncation=True, padding="longest", return_tensors="pt").to(device)
        
        summary_ids = model.generate(
            inputs["input_ids"],
            num_beams=beam_size,
            length_penalty=2.0,
            max_length=128,
            min_length=30,
            no_repeat_ngram_size=3,
            early_stopping=True
        )
        decoded = tokenizer.batch_decode(summary_ids, skip_special_tokens=True)
        summaries.extend(decoded)
        prog_bar.progress(min((i + batch_size) / len(chunks), 1.0))
        
    return summaries

# REPORT GENERATION
def create_pdf(summaries):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = styles['Title']
    body_style = styles['BodyText']
    
    story.append(Paragraph("ScholarScan Executive Summary", title_style))
    story.append(Spacer(1, 12))
    
    for i, summ in enumerate(summaries):
        header = Paragraph(f"<b>Section {i+1}</b>", styles['Heading2'])
        text = Paragraph(summ, body_style)
        story.append(header)
        story.append(text)
        story.append(Spacer(1, 12))
        
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_word(summaries):
    doc = Document()
    doc.add_heading('ScholarScan Executive Summary', 0)
    
    for i, summ in enumerate(summaries):
        doc.add_heading(f'Section {i+1}', level=1)
        doc.add_paragraph(summ)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# MAIN UI
st.title("📑 ScholarScan AI")
st.markdown("### Intelligent Document Summarization Pipeline")

# Load Model
tokenizer, model, device = load_model()

if not tokenizer:
    st.stop()

# Sidebar
with st.sidebar:
    st.header("⚙️ Settings")
    chunk_limit = st.slider("Chunk Size (Words)", 100, 300, 150, help="Smaller chunks = More detail.")
    beam_size = st.slider("Beam Size", 1, 5, 2, help="Higher = Better quality, Slower speed.")
    st.markdown("---")
    st.markdown("**Supported Formats:**\n- PDF\n- Word (.docx)\n- Images (.png, .jpg)\n- Text Input")

# Input Section
input_method = st.radio("Choose Input Method:", ["Upload File", "Type Text"], horizontal=True)
raw_text = ""

if input_method == "Upload File":
    uploaded_file = st.file_uploader("Upload your document", type=["pdf", "docx", "png", "jpg", "jpeg"])
    
    if uploaded_file:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == "pdf":
            with st.spinner("📄 Extracting text from PDF (OCR)..."):
                raw_text = extract_from_pdf(uploaded_file)
        elif file_type == "docx":
            raw_text = extract_from_docx(uploaded_file)
        elif file_type in ["png", "jpg", "jpeg"]:
            with st.spinner("🖼️ Extracting text from Image..."):
                raw_text = extract_from_image(uploaded_file)
                st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)

else:
    raw_text = st.text_area("Paste your text here:", height=300)

# Process Button
if raw_text and len(raw_text) > 10:
    st.info(f"✅ Text Detected: {len(raw_text.split())} words.")
    
    if st.button("🚀 Summarize Document"):
        with st.spinner("🤖 Generating Summary..."):
            final_summaries = summarize_text(raw_text, tokenizer, model, device, chunk_limit, beam_size)
            
        st.success("Summarization Complete!")
        
        # RESULTS SECTION
        col1, col2 = st.columns(2)
        
        # Column 1: Display Summary
        with col1:
            st.subheader("📝 Generated Summary")
            for i, summ in enumerate(final_summaries):
                with st.expander(f"Section {i+1}", expanded=True):
                    st.write(summ)
        
        # Column 2: Downloads
        with col2:
            st.subheader("💾 Download Report")
            
            # PDF Download
            pdf_file = create_pdf(final_summaries)
            st.download_button(
                label="📥 Download as PDF",
                data=pdf_file,
                file_name="Summary_Report.pdf",
                mime="application/pdf"
            )
            
            # Word Download
            word_file = create_word(final_summaries)
            st.download_button(
                label="📥 Download as Word (.docx)",
                data=word_file,
                file_name="Summary_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

elif raw_text:

    st.warning("⚠️ Text is too short to summarize (Min 10 words).")
